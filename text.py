from docx import Document

# Create a new Word Document
doc = Document()
doc.add_heading("PubMed Research Paper Fetcher - Full Report", 0)

# Summary
doc.add_heading("1. Summary", level=1)
doc.add_paragraph(
    "This project is a Python-based command-line tool that fetches research papers from the PubMed API "
    "based on a user-specified query. It filters out papers where at least one author is affiliated with a "
    "pharmaceutical or biotech company and outputs the results to a CSV file."
)

# Methodology
doc.add_heading("2. Methodology", level=1)

doc.add_heading("2.1 API Access", level=2)
doc.add_paragraph(
    "• Integrated PubMed's Entrez API using the Biopython library.\n"
    "• Supported full query syntax to enable precise and advanced searching.\n"
    "• Fetched XML summaries for the search results and parsed them for required data."
)

doc.add_heading("2.2 Affiliation Filtering", level=2)
doc.add_paragraph(
    "• Parsed the affiliation strings of each author.\n"
    "• Identified 'non-academic' authors using keywords such as 'Inc', 'Ltd', 'Biotech', 'Pharma', etc.\n"
    "• Papers with at least one such author were included in the final results."
)

doc.add_heading("2.3 Output Structure", level=2)
doc.add_paragraph(
    "• The CSV file contains the following columns:\n"
    "- PubmedID\n- Title\n- Publication Date\n"
    "- Non-academic Author(s)\n- Company Affiliation(s)\n- Corresponding Author Email"
)

# Implementation
doc.add_heading("3. Implementation", level=1)

doc.add_heading("3.1 Command-Line Features", level=2)
doc.add_paragraph(
    "• `--help` or `-h`: Show usage instructions\n"
    "• `--debug` or `-d`: Display debug output\n"
    "• `--file` or `-f [filename]`: Save output to the specified CSV file\n"
    "• Default behavior prints the output to the console if no file is given"
)

doc.add_heading("3.2 Poetry and Project Structure", level=2)
doc.add_paragraph(
    "• Used Poetry for virtual environment and dependency management.\n"
    "• Organized code into reusable modules and CLI entry point.\n"
    "• An executable script `get-papers-list` is created for streamlined use."
)

doc.add_heading("3.3 GitHub and Version Control", level=2)
doc.add_paragraph(
    "• Full project is version-controlled with Git.\n"
    "• Code is pushed to a GitHub repository as required."
)

# Results
doc.add_heading("4. Results", level=1)
doc.add_paragraph(
    "• Successfully ran multiple test queries like:\n"
    "  - 'covid AND vaccine AND 2023'\n"
    "  - 'oncology AND immunotherapy AND pharma'\n"
    "• The output CSV contains only those papers that meet the criteria.\n"
    "• Verified that email fields, affiliations, and publication dates were correctly captured."
)

# Evaluation Criteria
doc.add_heading("5. Evaluation Checklist", level=1)
table = doc.add_table(rows=1, cols=2)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Criteria'
hdr_cells[1].text = 'Status'

criteria = [
    ('Functional Requirements Met', '✅'),
    ('Full PubMed Query Syntax Supported', '✅'),
    ('Typed Python Used', '✅'),
    ('Readable Code with Docstrings', '✅'),
    ('Efficient API Usage', '✅'),
    ('Modular Structure', '✅'),
    ('Robust Error Handling', '✅'),
    ('Poetry for Dependency Setup', '✅'),
    ('Executable CLI via Poetry', '✅'),
    ('Documentation in README.md', '✅'),
    ('Published to GitHub', '✅'),
    ('Bonus: Split into Module + CLI', '✅'),
    ('Bonus: test.pypi Ready', '✅'),
]

for item, status in criteria:
    row = table.add_row().cells
    row[0].text = item
    row[1].text = status

# Notes
doc.add_heading("6. Notes", level=1)
doc.add_paragraph(
    "• All logic was implemented with typed Python for better maintainability.\n"
    "• Used Biopython and pandas for efficient parsing and data handling.\n"
    "• Added debugging and help support for better CLI usability.\n"
    "• Fully documented and structured as per task requirements."
)

# Save the document
full_report_path = "C:/Users/Rahul chaudhary/Desktop/PubMed_Fetcher_Full_Report.docx"
doc.save(full_report_path)

full_report_path
print("Report generated successfully at:", full_report_path)
print("✅ Report generated successfully.")
